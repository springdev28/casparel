#!/usr/bin/env python3
# @fileOverview: Repository tooling role: implements Build Casparel Reports for workspace development, build, validation, or documentation.
# System connection: invoked by package scripts or maintainers; it is not part of the end-user runtime bundle.
"""Build the branded Casparel product and Shipaton report artifacts."""

from __future__ import annotations

from datetime import date
from pathlib import Path
import textwrap
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DELIVERABLES = ROOT / "deliverables"
PDF_OUTPUT = ROOT / "output" / "pdf"
TMP = ROOT / "tmp" / "document_artifacts"
ASSETS = TMP / "assets"
LOGO_WIDE = ROOT / "artifacts" / "app" / "public" / "og-image.png"
LOGO_MARK = ROOT / "artifacts" / "app" / "public" / "apple-touch-icon.png"

PRODUCT_DOCX = DELIVERABLES / "Casparel-Product-Handbook-and-Audit.docx"
SHIPATON_DOCX = DELIVERABLES / "Casparel-Shipaton-2026-Readiness-and-Roadmap.docx"

NAVY = "13254A"
BLUE = "275DCE"
CYAN = "35B7EC"
INK = "17223B"
MUTED = "5B667A"
PALE_BLUE = "EAF1FF"
PALE_CYAN = "EAF9FE"
PALE_GRAY = "F3F5F8"
WHITE = "FFFFFF"
GREEN = "177A55"
PALE_GREEN = "E8F6F0"
GOLD = "A66A00"
PALE_GOLD = "FFF4DB"
RED = "A43B3B"
PALE_RED = "FBECEC"
TABLE_BORDER = "D8DFEA"

PAGE_WIDTH_DXA = 12240
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_row_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_borders(cell, color=TABLE_BORDER, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=TABLE_INDENT_DXA):
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if cell_mar is None:
        cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_mar)
    for side, value in (("top", 100), ("bottom", 100), ("start", 120), ("end", 120)):
        item = cell_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            cell_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        keep_table_row_together(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def set_paragraph_shading(paragraph, fill: str, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        borders = p_pr.find(qn("w:pBdr"))
        if borders is None:
            borders = OxmlElement("w:pBdr")
            p_pr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "8")
            el.set(qn("w:space"), "6")
            el.set(qn("w:color"), border_color)
            borders.append(el)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color_el, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.extend([r_pr, text_el])
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_numbering(doc: Document, *, bullet: bool, left: int, hanging: int) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "space")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.extend([tabs, ind])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, suffix, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def normalize_paragraph(paragraph, doc: Document):
    """Remove list/table inheritance that LibreOffice can render outside a cell."""
    paragraph.style = doc.styles["Normal"]
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:numPr", "w:tabs", "w:ind"):
        element = p_pr.find(qn(tag))
        if element is not None:
            p_pr.remove(element)
    paragraph.paragraph_format.left_indent = Inches(0)
    paragraph.paragraph_format.right_indent = Inches(0)
    paragraph.paragraph_format.first_line_indent = Inches(0)


def configure_running_header(header, subtitle: str):
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_after = Pt(0)


def configure_running_footer(footer):
    p = footer.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    add_page_number(p)


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def configure_document(doc: Document, *, title: str, subtitle: str, compact: bool):
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5 if compact else 11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20 if compact else 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 18 if compact else 16, 9 if compact else 8),
        "Heading 2": (13, BLUE, 13 if compact else 12, 6),
        "Heading 3": (11.5 if compact else 12, NAVY, 9 if compact else 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)

    for name, size in (("Title", 30), ("Subtitle", 14)):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(NAVY if name == "Title" else MUTED)
        style.paragraph_format.space_after = Pt(8)

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        configure_running_header(section.header, subtitle)
        configure_running_header(section.even_page_header, subtitle)

        section.footer.is_linked_to_previous = False
        section.even_page_footer.is_linked_to_previous = False
        configure_running_footer(section.footer)
        configure_running_footer(section.even_page_footer)

    props = doc.core_properties
    props.title = title
    props.subject = subtitle
    props.author = "Casparel"
    props.keywords = "Casparel, product documentation, audit, Shipaton 2026"
    props.comments = "Generated from the verified Casparel repository audit dated 15 August 2026."

    bullet = add_numbering(doc, bullet=True, left=540 if compact else 720, hanging=270 if compact else 360)
    decimal = add_numbering(doc, bullet=False, left=540 if compact else 720, hanging=270 if compact else 360)
    return bullet, decimal


def add_text(doc, text: str, *, bold_lead: str | None = None, italic=False, color=INK, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=10.5, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=10.5, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=10.5, color=color, italic=italic)
    return p


def add_bullets(doc, items: Iterable[str], num_id: int, *, compact=True):
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        p.paragraph_format.space_after = Pt(4 if compact else 6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        p.paragraph_format.line_spacing = Pt(14 if compact else 15)
        p.paragraph_format.keep_together = True
        run = p.add_run(item)
        set_run_font(run, size=10.25 if compact else 10.5, color=INK)


def add_numbered(doc, items: Iterable[str], num_id: int, *, compact=True):
    restarted_num_id = add_numbering(
        doc,
        bullet=False,
        left=540 if compact else 720,
        hanging=270 if compact else 360,
    )
    add_bullets(doc, items, restarted_num_id, compact=compact)


def add_heading(doc, text: str, level=1, *, font_size=None):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    if font_size is not None:
        for run in p.runs:
            set_run_font(run, size=font_size, color=BLUE if level < 3 else NAVY, bold=True)
    return p


def add_callout(doc, label: str, body: str, *, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    keep_table_row_together(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_borders(cell, accent, 8)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    lead = p.add_run(label.upper() + "  ")
    set_run_font(lead, size=9.5, color=accent, bold=True)
    text = p.add_run(body)
    set_run_font(text, size=10.25, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(doc, lines: Sequence[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, "EEF1F5", TABLE_BORDER)
    for idx, line in enumerate(lines):
        r = p.add_run(line)
        set_run_font(r, name="Courier New", size=8.5, color=NAVY)
        if idx != len(lines) - 1:
            r.add_break()


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], *, header_fill=PALE_BLUE, status_col=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        normalize_paragraph(p, doc)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) < 14 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(value)
        set_run_font(run, size=8.7, color=NAVY, bold=True)

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_borders(cell)
            if status_col is not None and idx == status_col:
                upper = value.upper()
                if any(k in upper for k in ("BLOCK", "FAIL", "NOT ", "P0", "LOW", "POOR", "HIGH")):
                    set_cell_shading(cell, PALE_RED)
                elif any(k in upper for k in ("PARTIAL", "MEDIUM", "P1", "CONDITIONAL", "DRAFT")):
                    set_cell_shading(cell, PALE_GOLD)
                elif any(k in upper for k in ("PASS", "READY", "COMPLETE", "STRONG", "FIXED")):
                    set_cell_shading(cell, PALE_GREEN)
            p = cell.paragraphs[0]
            normalize_paragraph(p, doc)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (
                idx == status_col or (idx == 0 and len(value) <= 25) or (idx > 0 and len(value) <= 7)
            ) else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(value)
            set_run_font(run, size=8.4, color=INK, bold=(idx == 0))
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_picture(doc, path: Path, *, width_inches: float, alt_text: str, caption: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    inline = run.add_picture(str(path), width=Inches(width_inches))
    doc_pr = inline._inline.docPr
    doc_pr.set("descr", alt_text)
    doc_pr.set("title", alt_text)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(0)
        cap.paragraph_format.space_after = Pt(8)
        r = cap.add_run(caption)
        set_run_font(r, size=8.5, color=MUTED, italic=True)


def add_section_break(doc):
    doc.add_page_break()


def add_contents(doc, items: Sequence[tuple[str, str]]):
    add_heading(doc, "Contents", 1)
    add_text(doc, "A navigational overview of this report. Section headings are preserved as real Word headings for the document outline and accessibility tools.", color=MUTED)
    rows = []
    for num, title in items:
        rows.append((num, title))
    add_table(doc, ("Section", "What it covers"), rows, (1150, 8210), header_fill=PALE_CYAN)


def cover_product(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(str(LOGO_WIDE), width=Inches(6.5))
    inline._inline.docPr.set("descr", "Casparel wordmark on a navy-to-blue background")
    inline._inline.docPr.set("title", "Casparel")

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Product Handbook & Technical Audit")
    set_run_font(r, size=28, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Product truth, architecture, operations, quality evidence, risks, and release guidance")
    set_run_font(r, size=13, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Verified 15 August 2026  |  Version 1.0  |  casparel.com")
    set_run_font(r, size=9.5, color=BLUE, bold=True)

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent_dxa=0)
    metrics = (("4", "CLIENT SURFACES"), ("173", "API TESTS"), ("39/39", "BROWSER RENDERS"), ("224 ms", "HTTP P95"))
    for idx, (value, label) in enumerate(metrics):
        cell = table.cell(0, idx)
        set_cell_shading(cell, NAVY if idx % 2 == 0 else BLUE)
        set_cell_borders(cell, WHITE, 4)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=16, color=WHITE, bold=True)
        r.add_break()
        r2 = p.add_run(label)
        set_run_font(r2, size=7.5, color=WHITE, bold=True)


def cover_shipaton(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(22)
    inline = p.add_run().add_picture(str(LOGO_MARK), width=Inches(0.95))
    inline._inline.docPr.set("descr", "Casparel brand mark")
    inline._inline.docPr.set("title", "Casparel")

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    r = kicker.add_run("EXECUTIVE COMPETITION REVIEW")
    set_run_font(r, size=9.5, color=CYAN, bold=True)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run("Shipaton 2026")
    set_run_font(r, size=34, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Readiness, winning strategy, submission plan, and future roadmap")
    set_run_font(r, size=14, color=MUTED)

    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2880, 6480], indent_dxa=0)
    score = table.cell(0, 0)
    set_cell_shading(score, NAVY)
    set_cell_borders(score, NAVY, 6)
    p = score.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("54/100")
    set_run_font(r, size=28, color=WHITE, bold=True)
    r.add_break()
    r2 = p.add_run("CURRENT READINESS")
    set_run_font(r2, size=8, color=CYAN, bold=True)

    verdict = table.cell(0, 1)
    set_cell_shading(verdict, PALE_BLUE)
    set_cell_borders(verdict, BLUE, 6)
    p = verdict.paragraphs[0]
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("Decision: enter only if store publication and transaction proof are completed by early September.")
    set_run_font(r, size=12, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Assessment date")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r.add_break()
    r2 = p.add_run("15 August 2026")
    set_run_font(r2, size=11, color=INK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Official deadline")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    r.add_break()
    r2 = p.add_run("30 September 2026, 11:45 PM PDT  |  1 October, 09:45 Istanbul")
    set_run_font(r2, size=11, color=INK)


FONT_REGULAR = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def chart_font(size: int, bold=False):
    return ImageFont.truetype(FONT_BOLD if bold else FONT_REGULAR, size)


def color(hex_value: str):
    return "#" + hex_value


def wrapped(draw: ImageDraw.ImageDraw, xy, text: str, width_chars: int, *, font, fill, spacing=8, anchor=None, align="left"):
    value = "\n".join(textwrap.wrap(text, width=width_chars))
    draw.multiline_text(xy, value, font=font, fill=fill, spacing=spacing, anchor=anchor, align=align)


def arrow(draw: ImageDraw.ImageDraw, start, end, *, fill=color(BLUE), width=5):
    draw.line([start, end], fill=fill, width=width)
    x2, y2 = end
    x1, y1 = start
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 > x1 else -1
        points = [(x2, y2), (x2 - direction * 18, y2 - 11), (x2 - direction * 18, y2 + 11)]
    else:
        direction = 1 if y2 > y1 else -1
        points = [(x2, y2), (x2 - 11, y2 - direction * 18), (x2 + 11, y2 - direction * 18)]
    draw.polygon(points, fill=fill)


def make_architecture_chart(path: Path):
    image = Image.new("RGB", (1800, 840), "white")
    draw = ImageDraw.Draw(image)
    title_font = chart_font(24, True)
    box_title = chart_font(25, True)
    body_font = chart_font(19)
    draw.text((900, 42), "CASPAREL SYSTEM MAP", font=title_font, fill=color(BLUE), anchor="mm")

    def box(x, y, w, h, title, detail, fill, edge=BLUE):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=color(fill), outline=color(edge), width=4)
        draw.text((x + 24, y + 22), title, font=box_title, fill=color(NAVY))
        draw.multiline_text((x + 24, y + 72), detail, font=body_font, fill=color(MUTED), spacing=8)

    box(45, 135, 350, 170, "Web", "React 19 + Vite\nComplete product surface", PALE_BLUE)
    box(45, 335, 350, 170, "Mobile", "Expo 54\nFocused native journeys", PALE_CYAN, CYAN)
    box(45, 535, 350, 170, "Desktop", "Electron shell\nControlled web navigation", PALE_GRAY, MUTED)
    box(650, 315, 485, 210, "Express API", "Auth · learning · resources · classes\nResearch · integrations · admin", PALE_BLUE)
    box(1390, 135, 365, 170, "PostgreSQL", "Drizzle schema\nOrdered migrations", PALE_GREEN, GREEN)
    box(1390, 335, 365, 170, "Research + Catalog", "OpenAI-compatible API\nOpen Library · Wikibooks", PALE_CYAN, CYAN)
    box(1390, 535, 365, 170, "External systems", "Google · RevenueCat\nHostinger", PALE_GOLD, GOLD)
    for sy in (220, 420, 620):
        arrow(draw, (395, sy), (645, 420))
    for ey in (220, 420, 620):
        arrow(draw, (1138, 420), (1385, ey))
    image.save(path)


def make_journey_chart(path: Path):
    image = Image.new("RGB", (1800, 570), "white")
    draw = ImageDraw.Draw(image)
    labels = [
        ("1", "DISCOVER", "Find open learning material"),
        ("2", "EVALUATE", "Check provenance and limitations"),
        ("3", "ORGANIZE", "Save to a list, class, or goal"),
        ("4", "STUDY", "Schedule and complete the next action"),
        ("5", "EVIDENCE", "Preserve work and learning signals"),
    ]
    fills = [PALE_BLUE, PALE_CYAN, PALE_GREEN, PALE_GOLD, PALE_BLUE]
    edges = [BLUE, CYAN, GREEN, GOLD, NAVY]
    for i, ((number, title, detail), fill, edge) in enumerate(zip(labels, fills, edges)):
        x = 40 + i * 355
        y, w, h = 110, 300, 340
        draw.rounded_rectangle((x, y, x + w, y + h), radius=20, fill=color(fill), outline=color(edge), width=4)
        draw.text((x + 24, y + 25), number, font=chart_font(27, True), fill=color(edge))
        draw.text((x + 24, y + 86), title, font=chart_font(24, True), fill=color(NAVY))
        wrapped(draw, (x + 24, y + 158), detail, 24, font=chart_font(19), fill=color(MUTED), spacing=9)
        if i < 4:
            arrow(draw, (x + w + 5, y + h // 2), (x + w + 48, y + h // 2), width=4)
    image.save(path)


def make_performance_chart(path: Path):
    image = Image.new("RGB", (1700, 760), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 50), "Production-safe smoke latency", font=chart_font(34, True), fill=color(NAVY))
    draw.text((70, 96), "20/20 HTTP checks · 0% failures · 1 virtual user", font=chart_font(21), fill=color(MUTED))
    labels = ["Overall", "Health", "Resources HTML", "Latest catalog", "Catalog search"]
    values = [224, 272, 188, 191, 202]
    fills = [BLUE, CYAN, NAVY, GREEN, GOLD]
    base_x, max_width = 390, 1080
    for idx, (label, value, fill) in enumerate(zip(labels, values, fills)):
        y = 185 + idx * 102
        draw.text((70, y + 22), label, font=chart_font(23, idx == 0), fill=color(INK))
        draw.rounded_rectangle((base_x, y, base_x + max_width, y + 64), radius=15, fill=color(PALE_GRAY))
        width = int(max_width * value / 320)
        draw.rounded_rectangle((base_x, y, base_x + width, y + 64), radius=15, fill=color(fill))
        draw.text((base_x + width + 18, y + 32), f"{value} ms", font=chart_font(22, True), fill=color(INK), anchor="lm")
    draw.text((base_x, 710), "p95 response time (milliseconds)", font=chart_font(19), fill=color(MUTED))
    image.save(path)


def make_readiness_chart(path: Path):
    image = Image.new("RGB", (1700, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "Current readiness profile · 54/100 overall", font=chart_font(34, True), fill=color(NAVY))
    labels = ["Product value", "Engineering", "Design + a11y", "RevenueCat", "Story + assets", "Mobile/store", "Traction"]
    values = [76, 72, 66, 48, 35, 30, 18]
    fills = [BLUE, BLUE, CYAN, GOLD, GOLD, RED, RED]
    base_x, max_width = 410, 1050
    for idx, (label, value, fill) in enumerate(zip(labels, values, fills)):
        y = 130 + idx * 101
        draw.text((70, y + 27), label, font=chart_font(23), fill=color(INK), anchor="lm")
        draw.rounded_rectangle((base_x, y, base_x + max_width, y + 56), radius=14, fill=color(PALE_GRAY))
        width = int(max_width * value / 100)
        draw.rounded_rectangle((base_x, y, base_x + width, y + 56), radius=14, fill=color(fill))
        draw.text((base_x + width + 18, y + 28), str(value), font=chart_font(23, True), fill=color(INK), anchor="lm")
    draw.text((base_x, 852), "Readiness score (0–100)", font=chart_font(19), fill=color(MUTED))
    image.save(path)


def make_timeline_chart(path: Path):
    image = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(image)
    draw.text((70, 45), "46-day critical path to submission", font=chart_font(34, True), fill=color(NAVY))
    phases = [
        ("Release candidate", 0, 7, BLUE),
        ("Prove journeys", 7, 10, CYAN),
        ("Stores + evidence", 17, 10, GREEN),
        ("Category proof", 27, 10, GOLD),
        ("Submission candidate", 37, 6, NAVY),
        ("Freeze", 43, 4, RED),
    ]
    base_x, max_width = 430, 1290
    for idx, (label, start, duration, fill) in enumerate(phases):
        y = 130 + idx * 88
        draw.text((70, y + 25), label, font=chart_font(22), fill=color(INK), anchor="lm")
        draw.rounded_rectangle((base_x, y, base_x + max_width, y + 52), radius=12, fill=color(PALE_GRAY))
        x1 = base_x + int(max_width * start / 47)
        x2 = base_x + int(max_width * (start + duration) / 47)
        draw.rounded_rectangle((x1, y, x2, y + 52), radius=12, fill=color(fill))
        draw.text(((x1 + x2) // 2, y + 26), f"{duration}d", font=chart_font(18, True), fill="white", anchor="mm")
    ticks = [(0, "15 Aug"), (7, "22 Aug"), (17, "1 Sep"), (27, "11 Sep"), (37, "21 Sep"), (43, "27 Sep"), (47, "1 Oct")]
    for day, label in ticks:
        x = base_x + int(max_width * day / 47)
        draw.line((x, 666, x, 681), fill=color(MUTED), width=2)
        draw.text((x, 700), label, font=chart_font(16), fill=color(MUTED), anchor="ma")
    image.save(path)


def build_charts():
    ASSETS.mkdir(parents=True, exist_ok=True)
    make_architecture_chart(ASSETS / "architecture.png")
    make_journey_chart(ASSETS / "journey.png")
    make_performance_chart(ASSETS / "performance.png")
    make_readiness_chart(ASSETS / "readiness.png")
    make_timeline_chart(ASSETS / "timeline.png")


def build_product_document():
    doc = Document()
    bullet, decimal = configure_document(
        doc,
        title="Casparel Product Handbook and Technical Audit",
        subtitle="Product Handbook & Audit",
        compact=True,
    )
    cover_product(doc)
    add_section_break(doc)

    add_heading(doc, "Executive brief", 1)
    add_callout(
        doc,
        "Product position",
        "Casparel is strongest when it moves a learner from discovery to source evaluation, organized study, and evidence. It should not position itself as a generic AI chatbot or social feed.",
        fill=PALE_CYAN,
        accent=CYAN,
    )
    add_text(
        doc,
        "The web/API foundation is materially stronger than a typical hackathon prototype: the production service is reachable, the automated API suite passes, a broad browser matrix is exercised, and teacher/resource workflows are unusually complete. The audit also found genuine release blockers in mobile packaging, quota behavior, product claims, public discovery, ownership controls, load-test safety, and coverage. Those findings were corrected rather than merely documented.",
    )
    add_table(
        doc,
        ("Area", "Assessment", "What matters now"),
        (
            ("Web + API", "Production-capable", "Keep monitoring; expand interaction and contract coverage."),
            ("Mobile", "Code-complete subset", "Prove signed builds, devices, stores, and purchases."),
            ("Desktop", "Controlled shell", "Signing, notarization, updater, and offline expectations remain."),
            ("Product evidence", "Synthetic evidence strong", "Run moderated student and teacher sessions."),
        ),
        (1500, 1950, 5910),
        status_col=1,
    )
    add_heading(doc, "What was verified", 2)
    add_bullets(
        doc,
        (
            "23 API test files and 173 tests passed in the isolated final run.",
            "The web production build completed across 2,734 transformed modules.",
            "The browser matrix passed 39 of 39 public/signed-in, theme, and viewport renders.",
            "Session handling passed valid, expired, and server-rejected token journeys.",
            "The desktop package compiled and passed four real Electron smoke scenarios.",
            "A production-safe HTTP smoke returned 20 of 20 HTTP 200 checks with 224 ms overall p95.",
        ),
        bullet,
    )
    add_callout(doc, "Evidence boundary", "No real account was created, no purchase was made, no recruited participant was tested, and no destructive production load was applied. Store distribution, RevenueCat dashboard state, and concurrent capacity remain unverified.", fill=PALE_GOLD, accent=GOLD)

    add_section_break(doc)
    add_contents(
        doc,
        (
            ("01", "Product definition and user value"),
            ("02", "User-facing surfaces and journeys"),
            ("03", "Architecture, repository, and request lifecycle"),
            ("04", "Identity, plans, limits, discovery, and research"),
            ("05", "Data model, integrations, and configuration"),
            ("06", "Development, deployment, and operations"),
            ("07", "Security, privacy, and accessibility"),
            ("08", "Quality, performance, and task-based user review"),
            ("09", "Implemented improvements"),
            ("10", "Risk register and prioritized next steps"),
            ("A", "Operator command reference and source register"),
        ),
    )

    add_section_break(doc)
    add_heading(doc, "01 · Product definition and user value", 1)
    add_text(doc, "Casparel is a cross-platform learning workspace for students and teachers. Students discover credible educational material, organize it around goals, schedule study, collaborate, and preserve evidence of progress. Teachers manage classes, recommend resources, create learning activities, and use evidence to support students.")
    add_picture(doc, ASSETS / "journey.png", width_inches=6.35, alt_text="Five-step Casparel value journey: discover, evaluate, organize, study, and evidence", caption="The product's durable differentiation is the complete credibility-to-action loop.")
    add_heading(doc, "Primary audiences", 2)
    add_table(
        doc,
        ("Audience / state", "Needs", "Casparel response"),
        (
            ("Signed-out visitor", "Understand value and browse safely", "Landing, public resource library, details, legal, and account entry."),
            ("Student", "Find trustworthy material and turn it into action", "Goals, schedules, resources, lists, classes, activities, canvases, collaboration, and evidence."),
            ("Teacher", "Guide classes and inspect learning", "Class management, assignments, recommendations, notes, seating, and Google Classroom integration."),
            ("Administrator", "Moderate and operate the system", "Verification queues, user/class controls, reports, and usage/cost visibility."),
            ("Premium account", "Use deep live-web research without account limits", "Unlimited account-level deep source research; server-wide safety budget still applies."),
        ),
        (1700, 2700, 4960),
    )
    add_section_break(doc)
    add_heading(doc, "Product promise and boundaries", 2)
    add_bullets(
        doc,
        (
            "The public learning-resource library and quick source check remain useful without Premium.",
            "Premium removes account-level limits from deep, cited, live-web source research; it does not remove the global emergency budget.",
            "Mobile is intentionally narrower than web and must never advertise missing native capabilities.",
            "Source reports are research aids, not objective trust scores; limitations and supporting links must remain visible.",
            "Desktop is a controlled online shell, not an offline-native product.",
        ),
        bullet,
    )

    add_heading(doc, "02 · User-facing surfaces and journeys", 1)
    add_heading(doc, "Web application", 2)
    add_table(
        doc,
        ("Area", "Representative pages", "Purpose"),
        (
            ("Public", "Landing, login, registration, legal, Resources, Support", "Acquisition, public discovery, trust, and account entry."),
            ("Home", "Dashboard, adaptive dashboard, guide", "Orientation, recommended action, and onboarding."),
            ("Learning", "Goals, activities, schedule, canvases", "Planning, focused work, and evidence."),
            ("Resources", "Search, detail, lists, catalog", "Discovery, evaluation, citation, and collections."),
            ("Community", "People, profiles, messages, forum", "Safe discovery and collaboration."),
            ("Teaching", "Classes and class detail", "Membership, assignments, resources, goals, notes, and seating."),
            ("Account", "Profile and settings", "Identity, integrations, privacy, and plan state."),
            ("Operations", "Admin", "Verification, moderation, support, and AI-usage visibility."),
        ),
        (1320, 3160, 4880),
    )
    add_heading(doc, "Mobile application", 2)
    add_text(doc, "The Expo Router app focuses on onboarding/login, dashboard, classes, schedule, resource browse/detail, quick and deep source review, profile, usage, and the Premium paywall. Native identifiers are com.casparel.app; the Expo slug and scheme are casparel. RevenueCat uses entitlement identifier premium and logs in with the numeric Casparel user ID.")
    add_callout(doc, "Truthful scope", "Goals, canvases, lists, messaging, forum, activities, admin, and AI-assisted catalog discovery are not complete native experiences. Paywall and store copy must describe only what the submitted binary actually provides.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "Desktop and shared design system", 2)
    add_text(doc, "The Electron package wraps CASPAREL_URL, validates navigation and deep links, preserves the window on embed failure, and has a real-window smoke test. The historical schoolar-edu package provides shared tokens and components across web and mobile; its legacy package name remains a compatibility detail.")
    add_heading(doc, "Critical journeys", 2)
    add_numbered(
        doc,
        (
            "Public discovery: open Resources, inspect starter material, search, and open a resource without registering.",
            "Credibility-to-action: run a quick or deep review, interpret strengths/limitations, then save or connect the resource to study work.",
            "Student organization: create a goal, plan time, complete activity, and preserve evidence.",
            "Teacher loop: create or manage a class, recommend/assign material, and review learning evidence.",
            "Subscription loop: use the free allowance, view the honest paywall at the effective limit, purchase/restore, and reconcile entitlement server-side.",
            "Session recovery: expire or reject a token, clear local state, and return safely to login.",
        ),
        decimal,
    )

    add_heading(doc, "03 · Architecture, repository, and request lifecycle", 1)
    add_picture(doc, ASSETS / "architecture.png", width_inches=6.35, alt_text="Architecture diagram connecting Casparel web, mobile, and desktop clients to the Express API, PostgreSQL, research catalogs, Google, RevenueCat, and hosting", caption="Client surfaces share one API and data model; desktop delegates to the web experience.")
    add_heading(doc, "Repository map", 2)
    add_table(
        doc,
        ("Path", "Responsibility"),
        (
            ("artifacts/api-server", "Express 5 API, authentication/authorization, integrations, and startup migrations."),
            ("artifacts/app", "React 19 + Vite SPA, production pages, and browser/session audit fixtures."),
            ("artifacts/mobile", "Expo 54 / React Native native client and RevenueCat adapter."),
            ("artifacts/desktop", "Electron shell, deep links, packaging, and runtime smoke."),
            ("artifacts/schoolar-edu", "Shared design tokens and web/native components."),
            ("lib/api-spec", "OpenAPI source of truth and code generation entrypoint."),
            ("lib/api-client-react", "Generated React Query client; never hand-edit generated files."),
            ("lib/api-zod", "Generated request/response schemas; never hand-edit generated files."),
            ("lib/db", "Drizzle schema and ordered migrations."),
            ("load-tests", "k6 smoke, public, and authenticated profiles with production guards."),
        ),
        (2500, 6860),
    )
    add_heading(doc, "Request lifecycle", 2)
    add_numbered(
        doc,
        (
            "A web or native client calls an endpoint below /api.",
            "The server applies security headers, CORS/body limits, and rate controls.",
            "Protected routes validate the signed bearer token; teacher/admin routes apply stronger authorization.",
            "Typed core routes validate payloads with generated Zod schemas.",
            "Data is read or written through Drizzle or parameterized SQL.",
            "Startup applies ordered migrations and checks database reachability/schema health before serving traffic.",
        ),
        decimal,
    )
    add_callout(doc, "Contract debt", "Forum, messaging, workflow, activity, canvas, expanded admin, and webhook surfaces are not all represented in OpenAPI. Migrate them spec-first in bounded domain batches to recover generated-client drift protection.", fill=PALE_RED, accent=RED)

    add_heading(doc, "04 · Identity, plans, limits, discovery, and research", 1)
    add_heading(doc, "Authentication and session behavior", 2)
    add_bullets(
        doc,
        (
            "Registration and login return a signed bearer token.",
            "The compatibility token key is schoolar_token on web and mobile; renaming it would break existing sessions.",
            "Expired or server-rejected tokens are cleared and redirect to login.",
            "Logout clears local session state; account deletion anonymizes/disables identity.",
            "The mobile user cache now uses casparel_user and migrates the legacy schooler_user value.",
        ),
        bullet,
    )
    add_heading(doc, "Plan reconciliation and quotas", 2)
    add_table(
        doc,
        ("Control", "Current rule", "Purpose"),
        (
            ("AI discovery fallback", "3 per user/day by default", "Contain optional fallback cost."),
            ("Deep research - day", "2 per free user/day", "Give every account a meaningful trial."),
            ("30-day research", "2 per free user/30-day window", "Longer-window cost control."),
            ("Global deep research", "20/day by default, including Premium", "Emergency server-wide budget."),
            ("Premium", "Account-level deep limits removed", "Monetized convenience; global safety remains."),
        ),
        (2200, 2860, 4300),
    )
    add_text(doc, "RevenueCat drives immediate native entitlement state and posts signed webhook events to POST /api/webhooks/revenuecat. The server records premium plan/expiry state. Persistent event idempotency and authoritative TRANSFER reconciliation remain release-critical work.")
    add_heading(doc, "Layered discovery", 2)
    add_numbered(
        doc,
        (
            "Query Casparel's stored public library.",
            "Query the locally stored open-education catalog.",
            "Fetch and cache Open Library or Wikibooks material when the first catalog page is sparse.",
            "Only when no catalog result exists and a feature flag permits it, use the paid AI fallback.",
        ),
        decimal,
    )
    add_heading(doc, "Source research", 2)
    add_text(doc, "Quick mode uses model knowledge. Deep mode performs live-web research and returns cited strengths, concerns, limitations, currency, and reputation analysis. Results are cached; concurrent work is restricted by user/resource; both account and global quotas apply. The interface must preserve uncertainty and supporting links.")

    add_heading(doc, "05 · Data model, integrations, and configuration", 1)
    add_heading(doc, "Data domains", 2)
    add_table(
        doc,
        ("Domain", "Representative records"),
        (
            ("Identity + safety", "Users, preferences, reports/safety state, activity log."),
            ("Teaching", "Classes, members, invitations, assignments, recommendations."),
            ("Learning", "Goals, evidence, workflow events, sessions, schedule blocks."),
            ("Resources", "Public/catalog resources, reviews, source-review cache, lists/items."),
            ("Collaboration", "Conversations, messages, forum content, comments/reports, canvases."),
            ("Integrations", "Google OAuth and Calendar tokens; subscription account state."),
        ),
        (2200, 7160),
    )
    add_callout(doc, "Migration invariant", "Every database schema change requires both the TypeScript schema edit and its generated Drizzle migration. Do not use drizzle push or direct ALTER TABLE. Timestamp fields feeding string schemas use mode: \"string\".", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "External integrations", 2)
    add_table(
        doc,
        ("Integration", "Purpose", "Required configuration"),
        (
            ("PostgreSQL", "Persistent state and quota counters", "DATABASE_URL"),
            ("OpenAI-compatible", "Discovery fallback and source research", "Base URL + API key"),
            ("Google Classroom", "Course import, students, resource sharing", "Client ID/secret + redirect"),
            ("Google Calendar / iCal", "Connected/exported scheduling", "OAuth settings + calendar redirect"),
            ("RevenueCat", "Products, offerings, entitlement", "Public SDK keys + webhook auth"),
            ("Hostinger", "Production hosting and smoke verification", "Deployment secrets + SITE_URL"),
        ),
        (1900, 3140, 4320),
    )
    add_heading(doc, "Required server configuration", 2)
    add_code_block(doc, ("DATABASE_URL=...", "SESSION_SECRET=...", "AI_INTEGRATIONS_OPENAI_BASE_URL=...", "AI_INTEGRATIONS_OPENAI_API_KEY=..."))
    add_text(doc, "Optional controls cover RevenueCat webhook authentication, Google OAuth, admin allowlists, CORS origins, catalog remote fetch, AI flags/quotas, logging, and strict database readiness. Secrets belong in host-managed environment variables or gitignored .env files, never source or committed documentation.")

    add_heading(doc, "06 · Development, deployment, and operations", 1)
    add_heading(doc, "Local startup", 2)
    add_code_block(doc, ("pnpm install --frozen-lockfile", "pnpm --filter @workspace/api-server run dev", "pnpm --filter @workspace/app run dev", "pnpm --filter @workspace/mobile run dev", "pnpm --filter @workspace/edu-ds run dev"))
    add_heading(doc, "Service boundaries", 2)
    add_table(
        doc,
        ("Service", "Default", "Rule"),
        (
            ("API", "8080 Replit / 5000 local", "Read PORT; server process remains separate."),
            ("Web", "23863", "Read PORT and API_URL; proxy /api to API."),
            ("Design system", "20495", "Run only for design-system development."),
            ("Mobile", "18115", "Expo development server; native store builds use EAS."),
        ),
        (1800, 2200, 5360),
    )
    add_heading(doc, "Quality gates", 2)
    add_code_block(doc, ("pnpm run typecheck", "pnpm --filter @workspace/api-server exec vitest run", "pnpm --filter @workspace/api-server run build", "pnpm --filter @workspace/app run build", "node artifacts/app/scripts/audit-pages.mjs", "node artifacts/app/scripts/audit-session.mjs", "pnpm --filter @workspace/desktop run compile", "pnpm --filter @workspace/desktop run smoke"))
    add_heading(doc, "Contract and schema workflow", 2)
    add_bullets(
        doc,
        (
            "API change: edit lib/api-spec/openapi.yaml, run the api-spec code generator, then add/register the Express handler.",
            "Generated React Query and Zod directories are output-only; never hand-edit them.",
            "Database change: edit the Drizzle schema, generate a migration, inspect and add safety guards, commit both together.",
            "The authentication token key and port/environment contracts are compatibility boundaries, not cosmetic naming opportunities.",
        ),
        bullet,
    )
    add_heading(doc, "Deployment and incident first checks", 2)
    add_numbered(
        doc,
        (
            "Call GET /api/healthz and inspect deployment logs.",
            "Confirm required environment variables are bound.",
            "Confirm migrations ran cleanly and the expected schema exists.",
            "Verify DATABASE_URL reachability and readiness diagnostics.",
            "Inspect RevenueCat webhook responses during entitlement incidents.",
            "Use only the bounded production smoke; run public/authenticated capacity profiles on staging with telemetry.",
        ),
        decimal,
    )

    add_heading(doc, "07 · Security, privacy, and accessibility", 1)
    add_heading(doc, "Implemented controls", 2)
    add_bullets(
        doc,
        (
            "Password hashing, signed tokens, route-level authorization, input/body limits, and request throttling.",
            "Parameterized data access, CORS policy, security headers, upload validation, and startup readiness checks.",
            "Moderation/reporting controls, minimal OAuth scope handling, signed RevenueCat webhooks, and redaction-aware logs.",
            "Public terms/privacy/support pages and an account deletion path.",
            "Automated browser assertions for runtime failures, contrast, accessible names, labels, headings, alt text, and overflow.",
        ),
        bullet,
    )
    add_heading(doc, "Residual security findings", 2)
    add_table(
        doc,
        ("Finding", "Severity", "Mitigation / resolution"),
        (
            ("RevenueCat TRANSFER reconciliation", "P1", "Persist event IDs and reconcile both subscriber aliases against authoritative state."),
            ("image-size ICNS denial of service", "High / build-time", "Avoid untrusted ICNS input; upgrade when a fixed release is available."),
            ("image-size JXL/HEIF denial of service", "High / build-time", "Avoid untrusted JXL/HEIF input; upgrade when a fixed release is available."),
            ("Newer routes outside OpenAPI", "P1", "Migrate domains spec-first and remove duplicate manual types."),
        ),
        (3500, 1900, 3960),
        status_col=1,
    )
    add_heading(doc, "Manual accessibility gates", 2)
    add_text(doc, "Automated checks cannot establish comprehension, focus usability, or screen-reader quality. Before native release, run VoiceOver and TalkBack on physical devices, Dynamic Type, reduced motion, keyboard/switch control, poor-network behavior, and package-selection/restore journeys.")

    add_section_break(doc)
    add_heading(doc, "08 · Quality, performance, and task-based user review", 1, font_size=13.5)
    add_heading(doc, "Verification dashboard", 2)
    add_table(
        doc,
        ("Check", "Result", "Evidence"),
        (
            ("API tests", "PASS", "23 files; 173 tests."),
            ("Web + mobile TypeScript", "PASS", "Workspace and focused final checks."),
            ("Expo public configuration", "PASS", "Casparel slug/scheme, IDs, production origin."),
            ("Web build", "PASS", "2,734 modules transformed."),
            ("Browser matrix", "PASS", "39/39 route/theme/viewport renders."),
            ("Session journeys", "PASS", "Valid, expired, and server-401 flows."),
            ("Desktop", "PASS", "Compile + four real-window smoke cases."),
            ("Production HTTP smoke", "PASS", "20/20 checks; 0% failures."),
            ("Dependency audit", "PARTIAL", "Three high findings fixed; two image-size highs remain."),
            ("Store install", "NOT VERIFIED", "Signed binaries/store credentials required."),
            ("Moderated usability", "NOT RUN", "Participant recruitment required."),
            ("Concurrent capacity", "NOT RUN", "Staging with database/CPU telemetry required."),
        ),
        (3050, 1760, 4550),
        status_col=1,
    )
    add_picture(doc, ASSETS / "performance.png", width_inches=6.15, alt_text="Horizontal bar chart of production-safe HTTP smoke p95 latency: overall 224 ms, health 272 ms, Resources 188 ms, latest catalog 191 ms, catalog search 202 ms", caption="Read-only, one-user latency sample; not a capacity claim.")
    add_heading(doc, "Web bundle profile", 2)
    add_table(
        doc,
        ("Chunk", "Raw", "Gzip", "Interpretation"),
        (
            ("Entry", "72.5 kB", "24.8 kB", "Healthy application entry."),
            ("React vendor", "185.8 kB", "58.6 kB", "Expected framework cost."),
            ("Radix vendor", "148.0 kB", "45.9 kB", "Shared accessible primitives."),
            ("Canvas", "203.6 kB", "65.9 kB", "Feature-specific; keep lazy."),
            ("Three.js", "617.4 kB", "157.0 kB", "Heavy but isolated."),
            ("p5", "1,129.2 kB", "331.0 kB", "Largest lazy dependency; measure before optimizing."),
        ),
        (1850, 1350, 1350, 4810),
    )
    add_heading(doc, "Task-based expert and synthetic review", 2)
    add_table(
        doc,
        ("Task", "Result", "Finding"),
        (
            ("Understand landing", "PASS", "Clear audiences and actions; no observed runtime errors."),
            ("Browse Resources signed out", "FIXED", "Starter cards now load from the existing top-rated API."),
            ("Clear search accessibly", "FIXED", "Explicit accessible name added."),
            ("Recover expired session", "PASS", "Token clears and login appears."),
            ("Remove own resource", "FIXED", "Remove control now appears only to the submitter."),
            ("Use free deep research", "FIXED", "Mobile consumes the server allowance before paywall."),
            ("Understand Premium", "FIXED", "Unsupported discovery/priority claims removed."),
            ("Select/restore accessibly", "PARTIAL", "Semantics improved; physical assistive-tech pass remains."),
            ("Teacher/admin workflow", "PARTIAL", "Major views render; live end-to-end creation/moderation was not executed."),
        ),
        (2830, 1500, 5030),
        status_col=1,
    )
    add_heading(doc, "Moderated research plan", 2)
    add_text(doc, "Recruit at least five students and five teachers, including keyboard-only and screen-reader representation where possible. Give tasks rather than tours; measure completion, time, misclicks, comprehension, confidence, and unaided recall.")
    add_numbered(
        doc,
        (
            "Find a credible resource for a real topic and explain why it is trustworthy.",
            "Save it, cite it, and connect it to a goal.",
            "Schedule study and find the plan again the next day.",
            "Join or manage a class and explain what is private versus shared.",
            "Compare quick and deep research and explain the Premium allowance.",
        ),
        decimal,
    )

    add_heading(doc, "09 · Implemented improvements", 1)
    improvements = (
        "Added mobile and api-spec to the pnpm workspace; regenerated a complete lockfile.",
        "Separated Expo runtime dependencies from build/type-only dependencies.",
        "Corrected Expo slug, scheme, and production Router origin to Casparel.",
        "Made root CI typecheck responsible for mobile instead of a redundant npx workaround.",
        "Let free mobile users consume their deep allowance before showing the paywall.",
        "Made usage reporting respect both day and 30-day deep counters.",
        "Removed unsupported Premium claims and strengthened native labels, roles, and selected state.",
        "Added public starter resources, public Support, and browser audit coverage.",
        "Restricted web resource removal controls to the submitting user.",
        "Corrected environment flags, budgets, and production origins.",
        "Protected casparel.com from accidental non-smoke k6 profiles and bounded smoke iterations.",
        "Made the Electron runtime smoke portable across Linux, macOS, and Windows layouts.",
        "Raised PostCSS and nanoid security floors, clearing three high-severity advisories.",
        "Added canonical product, audit, and Shipaton documentation.",
    )
    add_numbered(doc, improvements, decimal)

    add_section_break(doc)
    add_heading(doc, "10 · Risk register and prioritized next steps", 1)
    add_table(
        doc,
        ("Priority", "Risk", "Required action", "Exit evidence"),
        (
            ("P0", "No verified iOS/Android store release", "Configure EAS/credentials, build, install on devices, pass review.", "Public US listing + clean install."),
            ("P0", "RevenueCat lifecycle unproven", "Test purchase, renewal, restore, cancellation, expiry, transfer.", "Transaction matrix + server entitlement."),
            ("P1", "TRANSFER/idempotency gap", "Persist event IDs and reconcile authoritative subscriber state.", "Replay-safe tests and incident diagnostics."),
            ("P1", "No moderated user evidence", "Run five student + five teacher sessions and fix top failures.", "Task metrics + consented evidence."),
            ("P1", "Manual native accessibility incomplete", "VoiceOver, TalkBack, Dynamic Type, motion, keyboard/switch passes.", "Device checklist with no P0/P1 defects."),
            ("P1", "No client RUM/error funnel", "Add privacy-conscious telemetry and define SLOs.", "Activation/crash/performance dashboard."),
            ("P1", "Contract coverage incomplete", "Move newer API domains into OpenAPI in batches.", "Generated clients and drift tests."),
            ("P2", "Generic SPA metadata", "Add route-aware metadata or prerender/SSR.", "Crawl validation per public route."),
            ("P2", "Heavy specialist chunks", "Define route budgets and optimize from measured impact.", "RUM-backed performance regression budget."),
        ),
        (880, 2260, 3530, 2690),
        status_col=0,
    )
    add_callout(doc, "Release recommendation", "Web/API can continue in production with normal monitoring. Mobile is not yet store- or Shipaton-ready: signed builds, RevenueCat transaction proof, public publication, manual accessibility, and moderated usability are release gates.", fill=PALE_RED, accent=RED)

    add_section_break(doc)
    add_heading(doc, "Appendix A · Operator command reference", 1)
    add_heading(doc, "Install and validate", 2)
    add_code_block(doc, ("pnpm install --frozen-lockfile", "pnpm run typecheck", "pnpm --filter @workspace/api-server exec vitest run", "pnpm --filter @workspace/api-server run build", "pnpm --filter @workspace/app run build"))
    add_heading(doc, "Browser and desktop", 2)
    add_code_block(doc, ("node artifacts/app/scripts/audit-pages.mjs", "node artifacts/app/scripts/audit-session.mjs", "pnpm --filter @workspace/desktop run compile", "pnpm --filter @workspace/desktop run smoke"))
    add_heading(doc, "Production-safe smoke", 2)
    add_code_block(doc, ("BASE_URL=https://casparel.com SMOKE_ITERATIONS=5 pnpm run loadtest:smoke",))
    add_heading(doc, "Generate typed API clients", 2)
    add_code_block(doc, ("pnpm --filter @workspace/api-spec run codegen",))
    add_heading(doc, "Appendix B · Source register", 1)
    sources = (
        ("Casparel production", "https://casparel.com/"),
        ("RevenueCat Shipaton 2026 rules", "https://revenuecat-shipaton-2026.devpost.com/rules"),
        ("RevenueCat preparation guide", "https://revenuecat.github.io/codelabs/shipaton-2026-prep.html"),
        ("RevenueCat webhook event fields", "https://www.revenuecat.com/docs/integrations/webhooks/event-types-and-fields"),
        ("RevenueCat webhook event flows", "https://www.revenuecat.com/docs/integrations/webhooks/event-flows"),
        ("Expo Router production origin", "https://docs.expo.dev/versions/latest/sdk/router/"),
        ("image-size ICNS advisory", "https://github.com/advisories/GHSA-w3rx-r6r6-pgpr"),
        ("image-size JXL/HEIF advisory", "https://github.com/advisories/GHSA-5p2g-fcmc-qvqq"),
    )
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(label + ": ")
        set_run_font(r, size=9.5, color=INK, bold=True)
        add_hyperlink(p, url, url)
    add_text(doc, "Repository evidence is tied to commit e0dfa942d425f05fa4a65f439d0a890bbf1b4bb6. Competition rules and external service behavior can change; re-verify them before release or submission.", color=MUTED, italic=True)

    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc.save(PRODUCT_DOCX)


def build_shipaton_document():
    doc = Document()
    bullet, decimal = configure_document(
        doc,
        title="Casparel Shipaton 2026 Readiness and Roadmap",
        subtitle="Shipaton 2026 Readiness",
        compact=False,
    )
    cover_shipaton(doc)
    add_section_break(doc)

    add_heading(doc, "Executive decision", 1)
    add_callout(doc, "Verdict", "Casparel is a credible product, but it is not submission-ready today. Enter only if public store availability and RevenueCat transaction proof can be completed by early September.", fill=PALE_RED, accent=RED)
    add_text(doc, "The strongest assets are a working public product, an unusually complete education workflow, a real mobile codebase, RevenueCat integration code, and a memorable trust-before-study story. The largest blockers are evidence: no verified store-published binary, no demonstrated purchase lifecycle, no revenue or retention data, no two-minute demo, and no final submission asset pack.")
    add_table(
        doc,
        ("Decision lens", "Today", "Winning-state requirement"),
        (
            ("Eligibility", "Blocked / unverified", "Signed public iOS/Android app accessible in the US."),
            ("Monetization", "Partial", "Verified products, offering, purchase, restore, renewal, expiry, and transfer."),
            ("Category", "Design is best fit", "One sharp credibility-to-action experience with native polish."),
            ("Evidence", "Insufficient", "Real user outcomes, activation/retention/conversion, and first revenue."),
            ("Story", "Promising", "A sub-two-minute on-device narrative supported by proof."),
        ),
        (1900, 2320, 5140),
        status_col=1,
    )
    add_contents(
        doc,
        (
            ("01", "Official rule baseline and eligibility gates"),
            ("02", "Readiness score and probability bands"),
            ("03", "Category strategy and winning narrative"),
            ("04", "Two-minute demo plan"),
            ("05", "46-day execution roadmap"),
            ("06", "Evidence dashboard and submission kit"),
            ("07", "Post-Shipaton product roadmap"),
            ("A", "App-store/Devpost copy, assets, and final checklist"),
        ),
    )

    add_section_break(doc)
    add_heading(doc, "01 · Official rule baseline and eligibility gates", 1)
    add_text(doc, "The official rules require a qualifying native iOS, iPadOS, macOS, or Android app that integrates RevenueCat and is fully published by the deadline. The app must be accessible in the United States. Judges require free trial or promotional access when the experience is paid or gated.")
    add_table(
        doc,
        ("Official item", "Requirement used for this plan"),
        (
            ("Deadline", "30 September 2026, 11:45 PM PDT (1 October, 09:45 Istanbul)."),
            ("Native availability", "Fully published qualifying native app, accessible in the United States."),
            ("RevenueCat", "Integrated and demonstrably tied to the submitted paid/gated experience."),
            ("Demo", "Public YouTube or Vimeo video under two minutes."),
            ("Competition screenshot", "1179 × 2556 pixels, without a device frame."),
            ("Icon", "1024 × 1024 pixels."),
            ("Judge access", "Free trial, promo code, or equivalent access for gated content."),
            ("Next Gen exception", "Conditional; includes public open-source repository and other category-specific rules."),
        ),
        (2520, 6840),
    )
    add_callout(doc, "Rule caution", "Next Gen is not a shortcut for the current private repository. Use it only if the entrant independently qualifies, intentionally makes the project public under an appropriate license, and satisfies every category-specific condition.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "Gate checklist", 2)
    add_table(
        doc,
        ("Gate", "Status", "Required proof"),
        (
            ("Qualifying native platform", "PARTIAL", "Installable signed iOS/Android release candidate."),
            ("RevenueCat integration", "PARTIAL", "Dashboard offering, sandbox transaction, and entitlement evidence."),
            ("Fully published", "BLOCKED", "Public US store listing and downloadable build."),
            ("Judge access", "PARTIAL", "Clean reviewer account and trial/promo path."),
            ("Demo under two minutes", "NOT STARTED", "Public YouTube/Vimeo URL."),
            ("Screenshot and icon", "NOT COMPLETE", "Exact final exports with no device frame."),
            ("Submission evidence", "DRAFT", "Final text, URLs, screenshots, metrics, and claims."),
        ),
        (3000, 1800, 4560),
        status_col=1,
    )

    add_heading(doc, "02 · Readiness score and probability bands", 1, font_size=13.5)
    add_picture(doc, ASSETS / "readiness.png", width_inches=6.15, alt_text="Horizontal readiness chart showing product value 76, engineering 72, design and accessibility 66, RevenueCat 48, story and assets 35, mobile/store 30, and traction 18", caption="The 54/100 overall score is a prioritization model, not an official judging score.")
    add_heading(doc, "Winning probability bands", 2)
    add_table(
        doc,
        ("Scenario", "Judgment range", "Reason"),
        (
            ("Submit today", "<1%", "A missing publication gate can make the entry ineligible."),
            ("Complete minimum gates", "1–3%", "Functional, but weak differentiated evidence and traction."),
            ("Execute recommended plan", "5–10% category", "Polished native journey, transactions, demo, and real user evidence."),
            ("Grand Prize", "<2% without revenue", "The rules emphasize revenue shortlist and growth evidence."),
        ),
        (2450, 2200, 4710),
    )
    add_callout(doc, "Interpretation", "These percentages are judgment bands, not statistical forecasts: no entrant count, judge distribution, or historical base rate is available. The goal is to become installable, provable, understandable, and memorable in two minutes.", fill=PALE_CYAN, accent=CYAN)

    add_section_break(doc)
    add_heading(doc, "03 · Category strategy and winning narrative", 1)
    add_table(
        doc,
        ("Category", "Fit today", "Recommendation", "Proof needed"),
        (
            ("Design", "MEDIUM", "PRIMARY", "Native hierarchy, trust communication, accessibility, continuous screenshot story."),
            ("Replit Idea to Income", "CONDITIONAL", "SECONDARY IF PROVEN", "Required Replit workflow, three public posts, launch/growth/revenue evidence."),
            ("Peace", "MEDIUM-LOW", "EVIDENCE-DEPENDENT", "Measured information-literacy outcome; avoid overstated impact."),
            ("HAMM", "LOW", "DO NOT PRIORITIZE", "Conversion craft, pricing evidence, and diverse revenue streams."),
            ("Productivity Influencer", "POOR", "SKIP", "Product is not an Apple power-user content utility."),
            ("Next Gen", "CONDITIONAL", "ONLY IF ELIGIBLE", "Entrant criteria plus public open-source repo and license."),
        ),
        (1800, 1500, 2300, 3760),
        status_col=1,
    )
    add_heading(doc, "Recommended positioning", 2)
    add_callout(doc, "One-sentence pitch", "Casparel helps students move from ‘I found a resource’ to ‘I know why to trust it, what its limits are, and what to do next.’", fill=PALE_BLUE, accent=BLUE)
    add_text(doc, "Design is the strongest primary target because Casparel can make a complex credibility report immediately understandable to a student. The opportunity is not decorative polish: it is evidence, limitations, confidence, and next action arranged with exceptional clarity.")
    add_bullets(
        doc,
        (
            "Make the source-review reveal the visual center of the submitted app.",
            "Show evidence and limitations without fear-based or falsely objective scoring.",
            "Finish native accessibility, motion, empty/loading/error states, and package selection.",
            "Use screenshots that tell one continuous learner journey.",
            "Remove every claim that cannot be observed in the submitted binary.",
        ),
        bullet,
        compact=False,
    )

    add_heading(doc, "04 · Two-minute demo plan", 1)
    add_table(
        doc,
        ("Time", "Beat", "What the judge sees and understands"),
        (
            ("0:00–0:12", "Problem", "Search results are abundant; confidence is not."),
            ("0:12–0:35", "Discover", "Find a learning resource and show provenance."),
            ("0:35–1:05", "Evaluate", "Run deep source research; surface citations, concerns, and limitations."),
            ("1:05–1:25", "Act", "Save the resource and schedule the next study step."),
            ("1:25–1:42", "Teacher", "Recommend or assign it and expose learning evidence."),
            ("1:42–1:52", "Business", "Show the truthful paywall and a successful entitlement."),
            ("1:52–2:00", "Outcome", "Informed learning, not another AI answer box."),
        ),
        (1500, 1700, 6160),
    )
    add_callout(doc, "Creative rule", "Do not list every feature. The product breadth proves execution; the demo must make one credibility-to-action journey unforgettable.", fill=PALE_GOLD, accent=GOLD)

    add_heading(doc, "05 · 46-day execution roadmap", 1)
    add_picture(doc, ASSETS / "timeline.png", width_inches=6.3, alt_text="Timeline from 15 August to 1 October with six phases: release candidate, prove journeys, stores and evidence, category proof, submission candidate, and freeze", caption="Critical path from technical eligibility to a verifiable submission.")
    phases = (
        ("15–21 August · Release candidate", (
            "Require green CI on the audited baseline.",
            "Create or verify EAS project, Apple/Google identifiers, signing, and store accounts.",
            "Configure RevenueCat products, entitlement, offering, packages, public SDK keys, and webhook secret.",
            "Add persistent webhook idempotency and authoritative transfer reconciliation.",
            "Install internal iOS and Android builds on physical phones; freeze product truth and plan limits.",
        ), "A new reviewer can install, log in, restore, and see the same entitlement on client and server."),
        ("22–31 August · Prove critical journeys", (
            "Execute purchase, cancellation, renewal, billing issue, expiry, restore, and transfer tests.",
            "Run VoiceOver/TalkBack, Dynamic Type, reduced motion, keyboard, and poor-network checks.",
            "Recruit five students and five teachers for task-based sessions; fix the top breakdowns.",
            "Move the first newer API domain into OpenAPI; add privacy-conscious crash and funnel measurement.",
        ), "No P0/P1 defect in onboarding, research, purchase/restore, or account access."),
        ("1–10 September · Stores and evidence", (
            "Submit stable builds early to both stores and budget for rejection/remediation.",
            "Prepare reviewer notes, demo access, privacy forms, and support URLs.",
            "Instrument registration → resource → review → goal/save → paywall → purchase.",
            "Collect explicit, consented user quotes and trust-comprehension evidence.",
        ), "At least one store approved; the other has no known policy blocker; evidence grows daily."),
        ("11–20 September · Category proof", (
            "Lock the primary category and one honest secondary category.",
            "Produce screenshot narrative, icon, landing copy, and a 110–150 second rough demo.",
            "Show activation, return, research completion, paywall conversion, and first revenue.",
            "Complete any required public posts with useful build lessons and transparent metrics.",
        ), "Every judging claim has a URL, screenshot, metric, transaction, or user quote."),
        ("21–26 September · Submission candidate", (
            "Record the final demo below two minutes and export exact assets.",
            "Complete Devpost fields and category-specific answers.",
            "Test every public link signed out and from the US when possible.",
            "Run clean-install iOS/Android, production smoke, CI, and backup account tests.",
        ), "An unfamiliar reviewer can install, understand, and verify Casparel from the submission alone."),
        ("27–30 September · Freeze", (
            "No new features after 27 September; ship only release-blocking fixes with full regression.",
            "Re-check official rules and store availability; submit at least 24 hours early.",
            "Retain final media and answers offline; verify the saved submission after logout.",
        ), "Submission is live, independently verifiable, and backed up before the deadline."),
    )
    for title, actions, gate in phases:
        add_heading(doc, title, 2)
        add_bullets(doc, actions, bullet, compact=False)
        add_callout(doc, "Exit gate", gate, fill=PALE_GREEN, accent=GREEN)

    add_section_break(doc)
    add_heading(doc, "06 · Evidence dashboard and submission kit", 1)
    add_heading(doc, "Metrics to track daily", 2)
    add_table(
        doc,
        ("Metric", "Definition", "Direction"),
        (
            ("Activation", "New users completing a first source review in 24 hours", "UP"),
            ("Trust comprehension", "Users explaining two strengths and one limitation", "UP"),
            ("Time to value", "Registration to useful report", "DOWN"),
            ("Save-to-goal", "Reviewed resources connected to a goal", "UP"),
            ("D1 / D7 return", "Activated users returning", "UP"),
            ("Paywall conversion", "Unique viewers starting trial or purchasing", "UP"),
            ("Restore success", "Valid prior purchases restored", "100%"),
            ("Crash-free sessions", "Native sessions without fatal error", ">99.5%"),
            ("Research failure", "Deep requests ending in timeout/error/ambiguous limit", "DOWN"),
            ("Support response", "Time to resolve judge or user access issue", "DOWN"),
        ),
        (2200, 5500, 1660),
    )
    add_callout(doc, "Evidence rule", "Do not invent submission numbers. A small real cohort with precise definitions is more credible than a large unsupported claim.", fill=PALE_GOLD, accent=GOLD)
    add_heading(doc, "Submission asset set", 2)
    add_table(
        doc,
        ("Asset", "Specification", "Content"),
        (
            ("Icon", "1024 × 1024", "No alpha or pre-rounded corners."),
            ("Competition screenshot", "1179 × 2556", "No device frame; resolved UI state."),
            ("Store screenshots", "5–6 coherent frames", "Dashboard → Resources → research → Premium → schedule → plan."),
            ("Demo", "110–150 seconds; hard max under 2:00", "Real on-device discovery, evaluation, action, and transaction."),
            ("Judge access", "Stable account + trial/promo", "No dependence on founder intervention."),
            ("Proof pack", "Links + screenshots + metrics", "Every category claim backed by evidence."),
        ),
        (2200, 2850, 4310),
    )

    add_heading(doc, "07 · Post-Shipaton product roadmap", 1, font_size=13.5)
    add_table(
        doc,
        ("Horizon", "Objective", "Priority outcomes"),
        (
            ("0–3 months", "Reliability and learning loop", "Finish OpenAPI coverage; authoritative subscriptions; human-rated source research; goal-linked mobile actions; SLOs and retention cohorts."),
            ("3–6 months", "Teacher-led distribution", "Small class/school pilots; consent/admin controls; teacher credibility rubrics; collaborative evidence; institution-sponsored Premium test."),
            ("6–12 months", "Durable differentiation", "Provenance/evidence graph; curriculum alignment; privacy-preserving institution analytics; language/accessibility expansion; evaluate B2B licensing."),
        ),
        (1650, 2350, 5360),
    )
    add_callout(doc, "Strategic principle", "Build durable evidence infrastructure and teacher-led distribution only after student value, safety, subscription reliability, and trust comprehension are proven.", fill=PALE_CYAN, accent=CYAN)
    add_heading(doc, "Decision checkpoints", 2)
    add_bullets(
        doc,
        (
            "At three months, continue only the mobile and research bets that improve activation, trust comprehension, or subscription reliability.",
            "At six months, expand teacher distribution only when pilot classes show repeat use and privacy/consent operations are sustainable.",
            "At twelve months, evaluate school licensing only after learner value, safety, retention, and support economics are demonstrated.",
        ),
        bullet,
        compact=False,
    )
    add_heading(doc, "Roadmap guardrails", 2)
    add_bullets(
        doc,
        (
            "Do not chase broad mobile parity before the credibility-to-action loop is excellent.",
            "Do not describe AI outputs as objective truth; measure citation validity and human calibration.",
            "Do not add institution analytics until consent, minimization, export, and deletion controls are ready.",
        ),
        bullet,
        compact=False,
    )

    add_section_break(doc)
    add_heading(doc, "Appendix A · Ready-to-use submission copy", 1)
    add_heading(doc, "App identity", 2)
    add_table(
        doc,
        ("Field", "Recommended value"),
        (
            ("App name", "Casparel"),
            ("iOS subtitle", "Learn. Organize. Study."),
            ("Play short description", "Vetted open-education resources, classes, schedules, and AI research."),
            ("Category", "Education"),
            ("Bundle/package", "com.casparel.app"),
            ("Support", "https://casparel.com/support"),
            ("Marketing", "https://casparel.com"),
            ("Privacy", "https://casparel.com/privacy"),
        ),
        (2550, 6810),
    )
    add_heading(doc, "Store description", 2)
    for para in (
        "Casparel is where students and teachers find trustworthy learning materials and stay organized while they study.",
        "Discover open education. Browse a vetted catalog drawn from Open Library, Wikibooks, and other open sources. Every resource is free to open.",
        "Organize everything. Join classes, build reading lists, plan your schedule, and connect study work to the tools you already use.",
        "Research with AI. Casparel's Source Research helps explain who is behind a resource, what its strengths and limitations are, and how current the evidence appears. Use a quick check or request deep, cited live-web research.",
        "Casparel Premium. Upgrade for unlimited account-level deep source research. The core library stays free; Premium helps fund it.",
    ):
        add_text(doc, para)
    add_heading(doc, "Devpost tagline", 2)
    add_callout(doc, "Tagline", "A free, vetted open-education platform for students and teachers: organize your studies, evaluate any source with cited AI research, and keep the core library open with Premium.", fill=PALE_BLUE, accent=BLUE)
    add_heading(doc, "Screenshot narrative", 2)
    add_table(
        doc,
        ("#", "Screen", "Caption"),
        (
            ("1", "Dashboard", "Your studies, organized."),
            ("2", "Resources", "A vetted, free open-education library."),
            ("3", "Source Research", "Know who's behind any source."),
            ("4", "Paywall", "Go unlimited with Premium."),
            ("5", "Schedule", "Plan classes and study time."),
            ("6", "Profile / Plan", "Track your usage; upgrade anytime."),
        ),
        (650, 3060, 5650),
    )
    add_heading(doc, "Final submission checklist", 2)
    checklist = (
        "Signed iOS/Android release candidate installed and tested on physical devices.",
        "App publicly live and accessible in the United States before the deadline.",
        "RevenueCat products, premium entitlement, offering, and webhook live and mapped.",
        "Judge promo code or trial verified end-to-end.",
        "Purchase, restore, cancellation, expiry, renewal, billing issue, and transfer tested.",
        "1024 × 1024 icon and required 1179 × 2556 screenshot exported.",
        "Privacy policy and Support URLs live and accurate.",
        "Public demo under two minutes shows a real on-device transaction.",
        "Devpost categories selected only where every requirement is proven.",
        "Submission completed and verified at least 24 hours before the deadline.",
    )
    for item in checklist:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        p.paragraph_format.line_spacing = Pt(15)
        p.paragraph_format.keep_together = True
        box = p.add_run("☐  ")
        set_run_font(box, name="DejaVu Sans", size=10.5, color=BLUE)
        text = p.add_run(item)
        set_run_font(text, size=10.5, color=INK)

    add_heading(doc, "Appendix B · Sources and assumptions", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Official Shipaton 2026 rules: ")
    set_run_font(r, size=10, bold=True)
    add_hyperlink(p, "revenuecat-shipaton-2026.devpost.com/rules", "https://revenuecat-shipaton-2026.devpost.com/rules")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("RevenueCat submission preparation guide: ")
    set_run_font(r, size=10, bold=True)
    add_hyperlink(p, "revenuecat.github.io/codelabs/shipaton-2026-prep.html", "https://revenuecat.github.io/codelabs/shipaton-2026-prep.html")
    add_text(doc, "Readiness scoring and probability ranges are independent strategic judgments, not official RevenueCat scoring or statistical forecasts. Rules can change; repeat a full eligibility check immediately before submission.", color=MUTED, italic=True)
    add_callout(doc, "Final recommendation", "Choose Design as the product-quality target. Enter Replit Idea to Income only when the required workflow, public-post, growth, and revenue evidence are real. Spend the remaining runway on release proof, user evidence, and narrative discipline—not feature count.", fill=PALE_BLUE, accent=BLUE)

    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc.save(SHIPATON_DOCX)


def main():
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    PDF_OUTPUT.mkdir(parents=True, exist_ok=True)
    TMP.mkdir(parents=True, exist_ok=True)
    build_charts()
    build_product_document()
    build_shipaton_document()
    print(PRODUCT_DOCX)
    print(SHIPATON_DOCX)


if __name__ == "__main__":
    main()
